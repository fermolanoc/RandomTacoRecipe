"""
    Program that creates a workbook with 3 random taco recipes based
    on the recipe info by accessing an API that generates a random recipe each time it is accessed
    Workbook will include a Title with a related Taco image, credits to the image author, the taco API and
    workbook creator --> myself
    Then, 3 recipes will be shown on workbook, each one will have a name which will be composed based on the 5 main
    ingredients that the recipe contains (seasoning, condiment, mixin, base layer and shell).
    Each main ingredient will have its own title, description and detailed recipe info/preparation

    After each recipe document will include a page break to organize the content and show the next recipe
    By running the program it'll generate a Workbook document named RandomTacos.docx which will include 3 different
    taco recipes each time it is ran.
"""

import os
import requests
import docx
from PIL import Image, ImageDraw, ImageFont
import urllib.request
from docx.shared import Inches

# Getting image Full Size from Unsplash API
# Unsplash API to get images
api_url = 'https://api.unsplash.com/photos/random/'

# Developer Credentials - Access key
# In order for you to be able to run this code you should include your own Unsplash Access Key
# into your Environment variables
key = os.environ.get('IMAGE_KEY')

# Params to parse on the search
query = 'taco,tacos'

# get JSON data and convert it to Python dictionary by sending the api url, access key and search query on the request
image_data = api_url + '?query=' + query + '&client_id=' + key
random_image = requests.get(image_data).json()

# Get the full size image version
random_image_url = random_image['urls']['full']
# Save full size image into project folder
urllib.request.urlretrieve(random_image_url, 'random_taco.png')

# Resizing Image with Unsplash API -> dpr = Device Pixel Ratio -> Controls the density of the image, which means
# image density varies depending on the user's device so highest quality always will be shown depending on which quality
# level user's device can support -> min value is 1, max is 8 (3 is good enough and doesn't consume much bandwidth)
resized_image = random_image_url + '&w=600&dpr=3'

"""
    Another method to resized image manually is shown below
    I decided to go with the Unsplash API method above just because is faster, requires less code and
    final result has same great quality but both methods have been tested.
"""
# # Resizing image manually
# image = Image.open('random_taco.png')
# # RESIZING IMAGE
# # Getting original width and height
# width = image.width
# height = image.height
# print(width, height)
# percentage_to_reduce = float((800 * 100) / width)
# print(percentage_to_reduce)
#
# height_resize = int(height * (percentage_to_reduce/100))
# # width will be reduced by 14% too to keep contrast
# width_resize = int(width * (percentage_to_reduce/100))
#
# # Resize image keeping contrast
# resized_image_manually = image.resize((width_resize, height_resize))
# # print(image_resized.width, image_resized.height)
# resized_image_manually.show()
# # Save changes over image and include new width and height on file name with png format for better quality
# resized_image_manually.save(f'random_taco_{width_resize}x{height_resize}_manually.png')

# Save resized image into project folder
urllib.request.urlretrieve(resized_image, 'random_taco_600x600.png')

# DRAWING OVER IMAGE
# Open resized image
image = Image.open('random_taco_600x600.png')

# getting image size
height = image.height
W, H = (576, height)

# message to be printed over image
message = 'Random Taco Cookbook'

image_draw = ImageDraw.Draw(image)
# Write over image - font type and size
font = ImageFont.truetype('arial.ttf', 85)

# getting width and height of the message text
text_width, text_height = image_draw.textsize(message)
# text + coordinates where text will be placed + style
# W-w and H-h -> image width and height minus text width and height so text will be centered
image_draw.text(((W - text_width) / 2, (H - text_height) / 2), message, fill='purple', font=font)
image_draw.text(((W - text_width) / 2 + 5, (H - text_height) / 2 + 5), message, fill='yellow', font=font)

# Save edited image into project folder
image.save('random_taco_600x600_with_text.png')

# CREATE WORKBOOK
document = docx.Document()

# First Page includes Title, Image and Credits
# Title
document_title = 'Random Taco Cookbook'
document.add_paragraph(document_title.upper(), 'Title')

# Including image in Workbook
document.add_picture('random_taco_600x600_with_text.png', height=docx.shared.Inches(6.25))

# Add Credits
# Heading
credits_heading = 'Credits'
document.add_heading(credits_heading)

# Credits List
credit_p = document.add_paragraph
photo_author = random_image['user']['name']
credit_p(f'Taco image: Photo by {photo_author} on Unsplash', style='List Bullet')

# Random Taco Recipes's API URL
recipes_url = 'https://taco-1150.herokuapp.com/random/?full_taco=true'
credit_p(f'Tacos from: {recipes_url}', style='List Bullet')

# Code Author
credit_p('Code by: Fernando Molano', style='List Bullet')

document.add_page_break()
paragraph = document.add_paragraph()

# PASSING DATA TO WORKBOOK
# 5 main ingredients needed
ingredients = ['seasoning', 'condiment', 'mixin', 'base_layer', 'shell']

# empty list to store ingredients names, this will be used to create a custom title for each recipe
ingredients_list = []


# function that gets ingredients names, and their recipes based on url given when function is called
def create_recipe(url):
    # Get data in JSON format
    random_recipes_data = requests.get(url).json()
    # for each ingredient of the list above
    for ingredient in ingredients:
        # get ingredient name and save it into the ingredients_list
        ingredient_title = random_recipes_data[f'{ingredient}']['name']
        ingredients_list.append(ingredient_title)
    # call create_title function to create a custom title for the recipe
    create_title()
    # once custom title has been created and added to the workbook
    # for each ingredient in the list
    for ingredient in ingredients:
        # get recipe name and detail
        ingredient_recipe_name = random_recipes_data[f'{ingredient}']['name']
        ingredient_recipe = random_recipes_data[f'{ingredient}']['recipe']
        # print recipe name and detail on workbook
        document.add_heading(ingredient_recipe_name)
        document.add_paragraph(ingredient_recipe)


# function to create a custom title for the recipe using the ingredients names and adding some prepositions
# to create a nice sentence
def create_title():
    # get each item into the ingredients_list and add it to the sentence, at the end specify the text style = Title
    document.add_paragraph(ingredients_list[0] + ' with ' + ingredients_list[1] + ', '
                           + ingredients_list[2] + ' and '
                           + ingredients_list[3] + ' in ' + ingredients_list[4],
                           'Title')


# Code will be run 3 times
for i in range(3):
    # each time it'll call create_recipe function and send the url where Tacos Recipes data is stored
    create_recipe('https://taco-1150.herokuapp.com/random/?full_taco=true')
    # empty the list that was used to store ingredients names to create custom recipe's title so it can be use again
    ingredients_list = []
    # first 2 recipes add a page break after them so next recipe start in a new page
    if i != 2:
        document.add_page_break()
    # if it's the 3rd recipe just end the code so no blank page is added at the end
    else:
        break

# save the workbook as a .docx file into the project folder to save all the data that was added to it
document.save('Random_Taco_Recipes.docx')
